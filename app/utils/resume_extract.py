from __future__ import annotations

import io
from typing import Optional, Tuple


def _ext(filename: Optional[str]) -> str:
    if not filename:
        return ""
    name = filename.lower().strip()
    if "." not in name:
        return ""
    return name.rsplit(".", 1)[-1]


def extract_text_from_upload(filename: Optional[str], content_type: Optional[str], raw: bytes) -> Tuple[str, str]:
    """
    尝试从上传的简历文件中提取文本。

    返回：(text, method)
    method: txt|pdf|docx|unknown
    """
    ext = _ext(filename)
    ct = (content_type or "").lower()

    # 纯文本（txt/md）
    if ext in ("txt", "md", "text") or ct.startswith("text/"):
        try:
            return raw.decode("utf-8"), "txt"
        except Exception:
            # 常见情况：gbk
            try:
                return raw.decode("gbk"), "txt"
            except Exception:
                return "", "txt"

    # PDF
    if ext == "pdf" or ct in ("application/pdf",):
        # 优先：PyMuPDF（速度快、效果好）
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=raw, filetype="pdf")
            parts = []
            for page in doc:
                parts.append(page.get_text("text"))
            text = "\n".join(parts).strip()
            if text:
                return text, "pdf(pymupdf)"
        except Exception:
            pass

        # 兜底：pypdf（纯 Python，对部分 PDF 更稳）
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            parts = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
            text = "\n".join(parts).strip()
            if text:
                return text, "pdf(pypdf)"
        except Exception:
            pass

        # 多数扫描件 PDF 没有文本层，只能 OCR
        return "", "pdf(no_text_layer)"

    # DOCX
    if ext == "docx" or ct in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        try:
            from docx import Document  # python-docx
        except Exception:
            return "", "docx"
        try:
            bio = io.BytesIO(raw)
            doc = Document(bio)
            lines = []
            for p in doc.paragraphs:
                if p.text:
                    lines.append(p.text)
            # 也简单抓取表格文本
            for t in doc.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            text = "\n".join(lines).strip()
            return text, "docx"
        except Exception:
            return "", "docx"

    return "", "unknown"

